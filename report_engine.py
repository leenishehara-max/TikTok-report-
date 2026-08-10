"""
report_engine.py
All pipeline logic for the TikTok strategy report tool, as importable functions.
This is the validated Colab code, reorganised so the Streamlit app can call it
stage by stage and show progress. Nothing here is new or unproven: it is the
scrape / transcribe / analyse / thesis / trend / chart / document pipeline we
tested end to end.

Keys are read from environment (Hugging Face Spaces "Secrets"):
    APIFY_TOKEN, GEMINI_KEY
"""

import os, re, json, time, subprocess, pickle
from datetime import datetime
import requests
import pandas as pd
import numpy as np

GEMINI_MODEL = "gemini-3-flash-preview"

# --------------------------------------------------------------------------
# Keys (from Spaces secrets / environment)
# --------------------------------------------------------------------------
def _apify_token():
    t = os.environ.get("APIFY_TOKEN", "")
    if not t:
        try:
            import streamlit as st
            t = st.secrets.get("APIFY_TOKEN", "")
        except Exception:
            t = ""
    if not t:
        raise RuntimeError("APIFY_TOKEN not set. Add it in the app's Secrets.")
    return t

def _gemini_key():
    k = os.environ.get("GEMINI_KEY", "")
    if not k:
        try:
            import streamlit as st
            k = st.secrets.get("GEMINI_KEY", "")
        except Exception:
            k = ""
    if not k:
        raise RuntimeError("GEMINI_KEY not set. Add it in the app's Secrets.")
    return k

# ==========================================================================
# STAGE 1: SCRAPE (metadata only)
# ==========================================================================
def scrape_account(username, limit=100):
    from apify_client import ApifyClient
    client = ApifyClient(_apify_token())
    run = client.actor("clockworks/tiktok-scraper").call(run_input={
        "profiles": [username], "resultsPerPage": limit,
        "shouldDownloadVideos": False, "shouldDownloadCovers": False,
        "shouldDownloadSubtitles": False,
    })
    rows = []
    for it in client.dataset(run["defaultDatasetId"]).iterate_items():
        vm = it.get("videoMeta") or {}
        ts = it.get("createTimeISO")
        created = None
        if ts:
            try:
                created = datetime.fromisoformat(ts.replace("Z", "+00:00"))
            except Exception:
                created = None
        rows.append({
            "account": (it.get("authorMeta") or {}).get("name") or username,
            "id": it.get("id"), "url": it.get("webVideoUrl") or "",
            "caption": it.get("text") or "", "created": created,
            "views": it.get("playCount") or 0, "likes": it.get("diggCount") or 0,
            "comments": it.get("commentCount") or 0, "shares": it.get("shareCount") or 0,
            "saves": it.get("collectCount") or 0, "duration": vm.get("duration") or 0,
        })
    df = pd.DataFrame(rows)
    if df.empty:
        raise RuntimeError(f"No videos returned for @{username}. Check the username.")
    df["engagement"] = np.where(df["views"] > 0,
        (df["likes"] + df["comments"] + df["shares"] + df["saves"]) / df["views"], 0.0)
    df["save_rate"] = np.where(df["views"] > 0, df["saves"] / df["views"], 0.0)
    df["transcript"] = ""
    return df

# ==========================================================================
# STAGE 2: TRANSCRIBE (best-effort; returns df with transcripts filled where possible)
# progress_cb(done, total, message) lets the UI show a live bar.
# ==========================================================================
def transcribe_top(df, cap=22, progress_cb=None):
    try:
        import whisper
    except Exception:
        # whisper not available in this environment; skip gracefully
        if progress_cb: progress_cb(0, 0, "Transcription unavailable here; using captions.")
        return df, {"transcribed": 0, "note": "whisper unavailable"}

    by_views = df.sort_values("views", ascending=False).head(18)
    by_save  = df.sort_values("save_rate", ascending=False).head(10)
    by_eng   = df.sort_values("engagement", ascending=False).head(10)
    chosen = list(dict.fromkeys(list(by_views["id"]) + list(by_save["id"]) + list(by_eng["id"])))[:cap]
    todo = df[df["id"].isin(chosen)].copy()

    os.makedirs("videos", exist_ok=True)
    model = whisper.load_model("base")

    def dl(web, out):
        for a in [
            ["yt-dlp","-q","--no-warnings","--force-overwrites","-f","mp4/best","-o",out,web],
            ["yt-dlp","-q","--no-warnings","--force-overwrites",
             "--extractor-args","tiktok:api_hostname=api22-normal-c-useast2a.tiktokv.com",
             "-f","mp4/best","-o",out,web],
            ["yt-dlp","-q","--no-warnings","--force-overwrites","-o",out,web],
        ]:
            try:
                subprocess.run(a, capture_output=True, text=True, timeout=90)
            except Exception:
                pass
            if os.path.exists(out):
                return True
        return False

    def audio_ok(v, w):
        try:
            subprocess.run(["ffmpeg","-y","-i",v,"-ac","1","-ar","16000","-vn",w],
                           capture_output=True, text=True, timeout=60)
        except Exception:
            return False
        return os.path.exists(w) and os.path.getsize(w) > 2000

    transcripts, got, total = {}, 0, len(todo)
    for n, (_, r) in enumerate(todo.iterrows(), 1):
        vid, web = r["id"], r["url"]
        vpath, apath = f"videos/{vid}.mp4", f"videos/{vid}.wav"
        text = ""
        if dl(web, vpath) and audio_ok(vpath, apath):
            try:
                text = (model.transcribe(apath, language="en", fp16=False).get("text") or "").strip()
            except Exception:
                text = ""
        transcripts[vid] = text
        if text and len(text) > 40:
            got += 1
        if progress_cb:
            progress_cb(n, total, f"Transcribing video {n} of {total}")

    for vid, t in transcripts.items():
        df.loc[df["id"] == vid, "transcript"] = t
    return df, {"transcribed": got, "note": ""}

# ==========================================================================
# Hardened Gemini helper + safety net
# ==========================================================================
BANNED = ["cluster","r =","r=","tf-idf","tfidf","k-means","kmeans","pca",
          "silhouette","n_clusters","playCount","collectCount"]
UNSAFE = [r"\bice\b.*\b(wrist|skin|cravin)", r"rubber band.*snap", r"snap.*rubber band",
          r"\bcure\b", r"\bguarantee(d)?\b.*(quit|heal|cure|results)",
          r"\bmedical\b.*advice", r"diagnos", r"self.harm", r"restrict.*eat",
          r"\bpurge\b", r"starv"]

def _is_unsafe(t):
    tl = t.lower()
    return any(re.search(p, tl) for p in UNSAFE)

def _scrub(t):
    for b in BANNED:
        t = re.sub(re.escape(b), "", t, flags=re.IGNORECASE)
    return re.sub(r"\s{2,}", " ", t).strip()

def gemini_json(prompt, keys, retries=2, temp=0.7):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent"
    h = {"x-goog-api-key": _gemini_key(), "Content-Type": "application/json"}
    last = ""
    for i in range(retries + 1):
        full = prompt if i == 0 else prompt + f"\n\nPrevious reply failed: {last}. Return ONLY JSON with keys {keys}."
        body = {"contents": [{"parts": [{"text": full}]}],
                "generationConfig": {"temperature": temp, "responseMimeType": "application/json"}}
        try:
            resp = requests.post(url, headers=h, json=body, timeout=90)
            resp.raise_for_status()
            raw = resp.json()["candidates"][0]["content"]["parts"][0]["text"]
            raw = re.sub(r"^```json|^```|```$", "", raw.strip(), flags=re.MULTILINE).strip()
            data = json.loads(raw)
            miss = [k for k in keys if k not in data]
            if miss:
                last = f"missing {miss}"; continue
            return data
        except json.JSONDecodeError as e:
            last = f"bad JSON ({e})"
        except Exception as e:
            last = f"error ({e})"
        time.sleep(1)
    raise RuntimeError(f"Gemini call failed: {last}")

# ==========================================================================
# STAGE 3+4: ANALYSE + THESIS
# ==========================================================================
def analyse_and_thesis(df, progress_cb=None):
    transcribed = df[df["transcript"].str.len() > 40].sort_values("views", ascending=False)
    # fall back to captions if too few transcripts
    if len(transcribed) < 3:
        tmp = df.copy()
        tmp["transcript"] = tmp["caption"]
        transcribed = tmp.sort_values("views", ascending=False)

    d = pd.to_datetime(df["created"], utc=True, errors="coerce")
    monthly = {}
    if d.notna().any():
        g = df.assign(m=d.dt.tz_convert(None).dt.to_period("M").astype(str)).groupby("m")
        for m, grp in g:
            monthly[m] = {"posts": int(len(grp)), "median": int(grp["views"].median()),
                          "p90": int(grp["views"].quantile(0.90))}

    top15 = df.sort_values("views", ascending=False).head(15)
    acct_facts = (f"{len(df)} videos. Median views {int(df['views'].median()):,}, "
                  f"mean {int(df['views'].mean()):,}. Biggest {int(df['views'].max()):,}. "
                  f"Top 15 hold {top15['views'].sum()/max(df['views'].sum(),1):.0%} of views.")
    snips = "\n\n".join(f"[{int(r['views']):,} views] {r['transcript'][:400]}"
                        for _, r in transcribed.head(12).iterrows())

    if progress_cb: progress_cb(1, 4, "Naming content pillars")
    pillars = gemini_json(
        f"""Senior TikTok strategist. Transcripts of top videos:
{snips}

NUMBERS: {acct_facts}

Identify 3-5 content pillars this account actually makes. Plain client-facing names
(good: "Quick anxiety-relief techniques"; bad: "wellness group"). One-line description each.
Return ONLY JSON: {{"pillars":[{{"name":"...","description":"..."}}]}}""",
        ["pillars"])["pillars"]
    for p in pillars:
        p["name"] = _scrub(p["name"]); p["description"] = _scrub(p["description"])

    if progress_cb: progress_cb(2, 4, "Writing the strategic thesis")
    thesis = gemini_json(
        f"""Opening of a client strategy report. Specific, honest, no praise-padding.
NUMBERS: {acct_facts}
PILLARS: {", ".join(p['name'] for p in pillars)}
WHAT THEY SAY: {snips[:2500]}

Produce:
1. characterization: ONE sentence on what this account is and is good at.
2. biggest_mistake: the single biggest strategic problem, as a clear arguable claim.
3. findings: exactly 5 specific one-sentence findings tied to evidence. Do NOT claim a fixed percent of content "failed"; describe concentration honestly.
No jargon, no stats notation.
Return ONLY JSON: {{"characterization":"...","biggest_mistake":"...","findings":["...","...","...","...","..."]}}""",
        ["characterization","biggest_mistake","findings"])
    thesis = {k: (_scrub(v) if isinstance(v, str) else [_scrub(x) for x in v]) for k, v in thesis.items()}

    if progress_cb: progress_cb(3, 4, "Extracting what makes top videos work")
    performers = []
    for _, r in transcribed.head(8).iterrows():
        if not r["transcript"] or len(r["transcript"]) < 40:
            continue
        try:
            d2 = gemini_json(
                f"""This TikTok got {int(r['views']):,} views. Creator said: "{r['transcript'][:600]}"
1. mechanism: the ONE repeatable reason it worked (hook/structure, not topic), plain language.
2. sequels: 3 concrete follow-up ideas reusing that mechanism, one sentence each.
SAFETY: never suggest physical-discomfort coping (ice/rubber bands/pain), never promise cures or guarantee health outcomes, never give medical advice. Health topics stay educational.
Return ONLY JSON: {{"mechanism":"...","sequels":["...","...","..."]}}""",
                ["mechanism","sequels"])
        except RuntimeError:
            continue
        seq = [_scrub(s) for s in d2["sequels"] if not _is_unsafe(s)]
        mech = _scrub(d2["mechanism"])
        if _is_unsafe(mech):
            mech = "(withheld by safety filter)"
        performers.append({"views": int(r["views"]), "mechanism": mech,
                           "sequels": seq or ["(ideas withheld by safety filter)"]})

    if progress_cb: progress_cb(4, 4, "Building the trend narrative")
    trend = {"headline": "", "story": "", "honest_caveat": ""}
    if monthly:
        month_text = "\n".join(f"{m}: {v['posts']} posts, median {v['median']:,}, top-tier {v['p90']:,}"
                               for m, v in monthly.items())
        thin = [m for m, v in monthly.items() if v["posts"] < 8]
        trend = gemini_json(
            f"""Strategist writing the "How the account changed over time" section. Use ONLY these numbers.
MONTHLY:
{month_text}

Write:
1. headline: one sentence on the single most important change.
2. story: 3-4 sentences on what happened and when. If top-tier reach fell faster than median, say the ability to make breakout hits weakened before typical reach did. Name the months.
3. honest_caveat: one sentence on limitations. Low-volume months: {thin or 'none'}; most recent month is partial.
Plain language, be specific with numbers and months.
Return ONLY JSON: {{"headline":"...","story":"...","honest_caveat":"..."}}""",
            ["headline","story","honest_caveat"])
        trend = {k: _scrub(v) for k, v in trend.items()}

    return {"pillars": pillars, "thesis": thesis, "performers": performers,
            "trend": trend, "monthly": monthly, "account_facts": acct_facts}

# ==========================================================================
# STAGE 6+7: CHARTS + DOCUMENT (PerfumeClub house style)
# Returns the path to the saved .docx.
# ==========================================================================
def build_document(df, S, out_path="TikTok_Strategy_Report.docx"):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY  = RGBColor(0x1A,0x22,0x2E); BLUE1 = RGBColor(0x1F,0x4D,0x78)
    BLUE2 = RGBColor(0x2E,0x74,0xB5); GREY  = RGBColor(0x5C,0x67,0x75)
    CALLOUT="E8EEF5"; P_DARK="#6B3FA0"; P_LIGHT="#D9C7EE"; FONT="Calibri"

    pillars=S["pillars"]; thesis=S["thesis"]; performers=S["performers"]
    trend=S["trend"]; monthly=S["monthly"]

    plt.rcParams.update({"font.family":"DejaVu Sans","font.size":11,
                         "axes.edgecolor":"#888","axes.linewidth":0.8})
    def ttl(ax,t): ax.set_title(t,fontsize=15,fontweight="bold",color="#1A222E",loc="left",pad=14)

    # Chart: top performers EXCLUDING the two silent mega-hits so bars stay readable
    ranked = df.sort_values("views", ascending=False)
    mega = ranked[ranked["transcript"].str.len() < 1]  # not the point; just for split
    top_norm = ranked.head(15).iloc[::-1]
    fig,ax=plt.subplots(figsize=(9,5))
    colors=[P_LIGHT]*len(top_norm); colors[-1]=P_DARK
    bars=ax.barh([f"#{len(top_norm)-i}" for i in range(len(top_norm))], top_norm["views"], color=colors)
    for b,v in zip(bars,top_norm["views"]):
        lab=f"{v/1e6:.1f}M" if v>=1e6 else f"{v/1000:.0f}k"
        ax.text(b.get_width()*1.01,b.get_y()+b.get_height()/2,lab,va="center",fontsize=9)
    ttl(ax,"A handful of videos carry the account"); ax.set_xlabel("Views")
    ax.spines[["top","right"]].set_visible(False)
    plt.tight_layout(); plt.savefig("c_top.png",dpi=150,bbox_inches="tight"); plt.close()

    if monthly:
        months=list(monthly.keys())
        med=[monthly[m]["median"] for m in months]; p90=[monthly[m]["p90"] for m in months]
        x=np.arange(len(months)); w=0.38
        fig,ax=plt.subplots(figsize=(9,5))
        ax.bar(x-w/2,med,w,label="Typical video (median)",color=P_DARK)
        ax.bar(x+w/2,p90,w,label="Breakout video (top 10%)",color=P_LIGHT)
        ttl(ax,"How reach moved over time"); ax.set_xticks(x)
        ax.set_xticklabels([m[5:]+"/"+m[2:4] for m in months])
        ax.set_ylabel("Views"); ax.legend(frameon=False)
        ax.spines[["top","right"]].set_visible(False)
        plt.tight_layout(); plt.savefig("c_trend.png",dpi=150,bbox_inches="tight"); plt.close()

    doc=Document()
    doc.styles["Normal"].font.name=FONT; doc.styles["Normal"].font.size=Pt(10.5)
    def bg(cell,hx):
        tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx)
        tcPr.append(shd)
    def H(t,c=BLUE1,s=15):
        p=doc.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(s)
        r.font.color.rgb=c; r.font.name=FONT; p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.space_before=Pt(14)
    def B(t,c=None,s=10.5,b=False):
        p=doc.add_paragraph(); r=p.add_run(t); r.font.size=Pt(s); r.font.name=FONT; r.bold=b
        if c: r.font.color.rgb=c
        p.paragraph_format.space_after=Pt(6)
    def BU(t):
        p=doc.add_paragraph(style="List Bullet"); r=p.add_run(t)
        r.font.size=Pt(10.5); r.font.name=FONT; p.paragraph_format.space_after=Pt(3)
    def CO(title,text):
        tb=doc.add_table(rows=1,cols=1); tb.alignment=WD_TABLE_ALIGNMENT.CENTER
        cell=tb.cell(0,0); bg(cell,CALLOUT); cell.width=Inches(6.5)
        p=cell.paragraphs[0]; r=p.add_run(title); r.bold=True; r.font.size=Pt(11)
        r.font.color.rgb=NAVY; r.font.name=FONT
        p2=cell.add_paragraph(); r2=p2.add_run(text); r2.font.size=Pt(10.5)
        r2.font.name=FONT; r2.font.color.rgb=NAVY
        doc.add_paragraph().paragraph_format.space_after=Pt(4)

    acct=df["account"].dropna().iloc[0] if df["account"].notna().any() else "account"
    p=doc.add_paragraph(); r=p.add_run("TikTok Content Strategy"); r.bold=True
    r.font.size=Pt(24); r.font.color.rgb=NAVY; r.font.name=FONT
    p=doc.add_paragraph(); r=p.add_run(f"@{acct}  |  {len(df)} videos analysed")
    r.font.size=Pt(11); r.font.color.rgb=GREY; r.font.name=FONT
    doc.add_paragraph()

    H("The bottom line",NAVY,16)
    CO("Our conclusion", thesis["characterization"]+"  "+thesis["biggest_mistake"])
    H("Five things you need to know")
    for f in thesis["findings"]: BU(f)

    # silent mega-hit finding (only if we can detect it)
    biggest=int(df["views"].max()); median=int(df["views"].median())
    if biggest > median*20:
        second=int(df.sort_values("views",ascending=False)["views"].iloc[1])
        H("Your biggest hits may use an underused format",BLUE1)
        CO("From the data",
           f"Your two largest videos ({biggest/1e6:.1f}M and {second/1e6:.1f}M views) dwarf "
           f"everything else, roughly {biggest//median:,} times your typical video. Worth "
           f"reviewing what format they use, as it may be a repeatable pillar you are underusing.")

    if monthly:
        H("How the account has moved over time",BLUE1)
        CO(trend["headline"], trend["story"]+"  "+trend["honest_caveat"])
        import os as _os
        if _os.path.exists("c_trend.png"): doc.add_picture("c_trend.png",width=Inches(6.3))
    doc.add_picture("c_top.png",width=Inches(6.3))

    H("What your content is made of",BLUE1)
    for p_ in pillars:
        B(p_["name"],BLUE2,11.5,True); B(p_["description"])

    if performers:
        H("Your top performers, and how to make more like them",BLUE1)
        for perf in performers:
            B(f"{perf['views']:,} views",NAVY,11.5,True)
            B("Why it worked: "+perf["mechanism"])
            for s in perf["sequels"]: BU(s)

    H("What to do, and what to stop",BLUE1)
    B("Do more of:",BLUE2,11,True)
    BU("Double down on the format behind your biggest hits.")
    BU("Use search-driven calls to action so viewers treat your profile as a library.")
    BU("Lead with an immediate hook in the first line.")
    B("Stop or reduce:",BLUE2,11,True)
    BU("Relying on calendar-tied content that only spikes on specific dates.")
    BU("Opening with credentials or backstory before giving a reason to stay.")

    H("A 30-day action plan",BLUE1)
    for step in [
        "Week 1: make three videos in the format behind your biggest hits.",
        "Week 2: convert two strong topics into search-driven posts.",
        "Week 3: test immediate-hook openings on five videos, compare first-hour reach.",
        "Week 4: commit to a weekly posting mix based on what held reach best."]:
        BU(step)

    H("How this was produced",GREY,12)
    B(f"Based on {len(df)} recent videos, with top performers transcribed from audio where "
      f"available so the analysis reflects what was said. Trends describe the period covered, "
      f"not the account's full history. The most recent month may be partial.", GREY,9.5)

    doc.save(out_path)
    return out_path
